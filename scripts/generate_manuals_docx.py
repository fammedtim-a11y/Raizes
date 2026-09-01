from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "assets" / "logo-raizes-kids.png"
OUT_PROGRAMADOR = ROOT / "Manual do Programador.docx"
OUT_USUARIO = ROOT / "Manual do Usuario.docx"
UPDATED_AT = "27/08/2026"

BLUE = "2D9CFF"
GREEN = "6BCB3D"
YELLOW = "FFD93D"
RED = "FF5A5A"
PURPLE = "7B61FF"
TURQUOISE = "29C7C9"
DARK = "17324D"
MUTED = "5D6B7A"
LIGHT = "E8EEF5"


def rgb(hex_color):
    hex_color = hex_color.strip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = rgb(color)


def setup_doc(title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    doc.core_properties.title = title
    doc.core_properties.author = "Raízes Kids"
    return doc


def add_footer(doc, label):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = f"{label} | Raízes Kids | atualizado em {UPDATED_AT}"
    paragraph.runs[0].font.size = Pt(8)
    paragraph.runs[0].font.color.rgb = rgb(MUTED)


def add_cover(doc, title, subtitle, audience, accent=BLUE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(2.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run(run, 26, True, accent)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run(run, 13, False, DARK)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, "F4F8FC")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(audience)
    set_run(run, 11, True, DARK)
    doc.add_paragraph("Documento de referência do sistema Raízes Kids. Use este material para desenvolvimento, operação, treinamento e repasse de conhecimento.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text):
    return doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        shade(hdr[idx], LIGHT)
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                set_run(r, 10, True, DARK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    set_run(r, 9.5, False, DARK)
    if widths is None:
        widths = [int(9360 / len(headers))] * len(headers)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, title, text, fill="F4F8FC", accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, 11, True, accent)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def build_programmer_manual():
    doc = setup_doc("Manual do Programador")
    add_footer(doc, "Manual do Programador")
    add_cover(
        doc,
        "Manual do Programador",
        "Documentação técnica do sistema Raízes Kids",
        "Público: desenvolvedores, suporte técnico, integradores e responsáveis por manutenção",
        BLUE,
    )

    add_h1(doc, "1. Visão Geral")
    add_para(doc, "O Raízes Kids é uma aplicação web para ministério com crianças. O sistema organiza lições bíblicas por faixa etária, trilhas com vídeos do YouTube, cultos em família, treinamentos, EBF, novidades, usuários, controle de acesso e comunicação com a base de membros.")
    add_callout(doc, "Resumo técnico", "Aplicação Node.js sem framework web externo, front-end em HTML/CSS/JavaScript puro, persistência em arquivos JSON no diretório configurado por DATA_DIR e publicação prevista em Render conectada ao GitHub.", "EEF9FF", BLUE)

    add_h2(doc, "Objetivos do sistema")
    add_bullets(doc, [
        "Centralizar conteúdos bíblicos infantis em um catálogo filtrável por idade, categoria, testamento, data especial e mês de inclusão.",
        "Permitir que administradores cadastrem, editem e removam conteúdos sem alterar código.",
        "Disponibilizar leitura na própria página, exportação de PDF e ebook, além de trilhas em vídeo com experiência visual de streaming.",
        "Controlar usuários, aprovação, nível de acesso, licença e registros de acesso.",
        "Apoiar comunicação contínua com usuários por links de email e WhatsApp."
    ])

    add_h1(doc, "2. Arquitetura")
    add_table(doc, ["Camada", "Responsabilidade", "Arquivos principais"], [
        ["Servidor", "Rotas HTTP, autenticação, segurança, persistência JSON, uploads e backup.", "server.js"],
        ["Interface pública", "Home, filtros, catálogo, leitor de lições, trilhas, novidades, rodapé e módulos de conteúdo.", "index.html, app.js, styles.css"],
        ["Gerenciamento", "Cadastros de lições, trilhas, novidades, EBF, treinamentos, usuários, contato e comunicação.", "gerenciamento.html, app.js, auth.js"],
        ["Autenticação", "Login, cadastro, sessão, perfil, usuários, logs e ações administrativas.", "login.html, perfil.html, auth.js"],
        ["Dados iniciais", "Seeds/importações usadas para popular conteúdos quando os JSON ainda não existem.", "lessons-data.js, seed-content.js, seed-imports.js"],
        ["Facilitador desktop", "Aplicação Electron auxiliar empacotável.", "facilitador/"],
    ], [1800, 4300, 3260])

    add_h2(doc, "Persistência e diretórios")
    add_para(doc, "O servidor usa DATA_DIR para decidir onde salvar os dados. Em ambiente local, o padrão é a pasta data dentro do projeto. Em produção no Render, recomenda-se usar disco persistente e configurar DATA_DIR para o caminho do disco.")
    add_table(doc, ["Arquivo", "Conteúdo"], [
        ["users.json", "Usuários, papéis, aprovação, status, nível de acesso, licença e hashes de senha."],
        ["lessons.json", "Lições bíblicas cadastradas e imagens persistidas."],
        ["videos.json", "Vídeos manuais das Trilhas."],
        ["devotionals.json", "Cultos em Família."],
        ["trainings.json", "Treinamentos e anexos."],
        ["ebf.json", "Materiais de EBF."],
        ["notifications.json", "Novidades exibidas no sino, home e painel."],
        ["communications.json", "Campanhas criadas pelo administrador master."],
        ["sessions.json", "Sessões persistidas."],
        ["access-log.json", "Histórico de acessos e dispositivos."],
        ["site-info.json", "Sobre, email, WhatsApp, Instagram e URL do site."],
        ["uploads/", "Arquivos enviados: imagens, anexos e materiais."],
    ], [2400, 6960])

    add_h1(doc, "3. Módulos Funcionais")
    modules = [
        ("Home", "Apresenta a marca, chamadas principais, faixas etárias, novidades, login/cadastro e links institucionais."),
        ("Catálogo de Lições", "Exibe cards de lições, leitor com seções pedagógicas, filtros, imagens, atividade de colorir e exportação PDF da lição."),
        ("Ebook", "Imprime todas as lições ativas no filtro, usando layout próprio de impressão e marca autoral."),
        ("Trilhas", "Reúne vídeos manuais e vídeos encontrados nos textos das lições/conteúdos. Possui player integrado, destaque principal e prateleiras por categoria."),
        ("Culto em Família", "Conteúdo devocional com princípio, texto bíblico, oração, atividade e apoio visual."),
        ("Treinamentos", "Materiais para líderes, com vídeo, imagem, texto, orientações e anexos."),
        ("EBF Completa", "Cadastro e exibição de materiais de Escola Bíblica de Férias, programação, orientações e anexos."),
        ("Novidades", "Central de chamadas para conduzir o usuário a novos conteúdos."),
        ("Comunicação", "Área exclusiva do administrador master para criar campanhas por público e gerar links de email/WhatsApp."),
        ("Usuários e Acesso", "Aprovação de cadastros, bloqueio, reativação, senha, categoria de acesso, licença e logs."),
    ]
    add_table(doc, ["Módulo", "Descrição"], modules, [2100, 7260])

    add_h1(doc, "4. Cadastros e Estrutura de Dados")
    add_h2(doc, "Lições")
    add_para(doc, "A lição é o conteúdo central do sistema. Cada registro possui metadados, imagens opcionais e um objeto sections com os blocos pedagógicos.")
    add_table(doc, ["Campo", "Descrição"], [
        ["id", "Identificador único."],
        ["title", "Nome da lição."],
        ["category", "Tema bíblico, como Deus Pai, Criação, Fé ou Gratidão."],
        ["age", "Faixa etária normalizada."],
        ["verse", "Versículo principal."],
        ["createdAt", "Data de inclusão usada no filtro mensal."],
        ["cardImage", "Imagem de capa do card."],
        ["activityImage", "Imagem da atividade de colorir."],
        ["sections.objectives", "Objetivos da lição."],
        ["sections.welcome", "Recepção e acolhimento."],
        ["sections.icebreaker", "Quebra-gelo."],
        ["sections.openingPrayer", "Oração inicial."],
        ["sections.worshipOffering", "Louvor e oferta."],
        ["sections.bibleLesson", "Lição bíblica."],
        ["sections.practice", "Aplicação prática."],
        ["sections.memoryVerse", "Memorização do versículo."],
        ["sections.activity", "Atividade."],
        ["sections.finalPrayer", "Oração final."],
        ["sections.snack", "Lanche."],
    ], [2700, 6660])

    add_h2(doc, "Vídeos / Trilhas")
    add_table(doc, ["Campo", "Descrição"], [
        ["title", "Título do vídeo."],
        ["url / youtubeId", "Link do YouTube e ID extraído para player/thumbnail."],
        ["category", "Categoria da trilha."],
        ["age", "Faixa etária opcional."],
        ["lessonId", "Lição relacionada quando houver."],
        ["description", "Observação de uso pedagógico."],
        ["playlist / season", "Organização editorial."],
        ["featured / trending / recommended", "Controle de destaque e ordenação."],
    ], [2700, 6660])

    add_h2(doc, "Usuários")
    add_table(doc, ["Campo", "Descrição"], [
        ["username", "CPF ou identificador usado no login."],
        ["passwordHash", "Senha gravada com hash scrypt."],
        ["role", "admin ou user."],
        ["accessLevel", "simple, test, leader ou prime."],
        ["approved", "Indica se o administrador aprovou o cadastro."],
        ["active", "Permite bloquear ou reativar acesso."],
        ["licenseExpiresAt", "Data de vencimento do acesso."],
        ["name, email, phone, address, church, churchCity", "Dados cadastrais do usuário."],
        ["lastLoginAt / lastAccessAt", "Dados operacionais para auditoria."],
    ], [3000, 6360])

    add_h1(doc, "5. Rotas Principais da API")
    add_table(doc, ["Rota", "Método", "Uso"], [
        ["/api/session", "GET", "Retorna o usuário logado e mensagens de sessão."],
        ["/api/login", "POST", "Autentica usuário."],
        ["/api/logout", "POST", "Encerra sessão."],
        ["/api/register", "POST", "Cria cadastro aguardando aprovação."],
        ["/api/password-reset", "POST", "Solicitação/redefinição simples de senha."],
        ["/api/profile", "GET/POST", "Consulta e edita perfil do usuário logado."],
        ["/api/lessons", "GET", "Lista lições para o catálogo público."],
        ["/api/videos", "GET", "Lista vídeos manuais."],
        ["/api/novidades", "GET", "Lista novidades ativas."],
        ["/api/admin/lessons", "GET/POST", "Lista/salva lições. Requer admin."],
        ["/api/admin/videos", "GET/POST", "Lista/salva trilhas manuais. Requer admin."],
        ["/api/admin/novidades", "GET/POST", "Lista/salva novidades. Requer admin."],
        ["/api/admin/users", "GET", "Lista usuários para gestão. Requer admin."],
        ["/api/admin/comunicacao/*", "GET/POST", "Audiência e campanhas. Requer administrador master."],
        ["/api/admin/backup", "GET", "Baixa backup geral tar.gz. Requer admin."],
    ], [3300, 1300, 4760])

    add_h1(doc, "6. Segurança")
    add_bullets(doc, [
        "As páginas administrativas são protegidas no servidor; gerenciamento.html exige usuário admin.",
        "A autenticação usa cookie de sessão HttpOnly e SameSite=Lax.",
        "Senhas são armazenadas com hash scrypt.",
        "Há limitação básica de tentativas de login por IP/usuário.",
        "O servidor aplica cabeçalhos de segurança, incluindo Content-Security-Policy, X-Content-Type-Options e proteção contra clickjacking.",
        "A pasta data não deve ser servida como estático; o acesso passa pelas rotas controladas.",
        "A Comunicação fica restrita ao usuário master 08047232657.",
        "Logs de acesso registram evento, rota, usuário, IP e dispositivo para auditoria."
    ])
    add_callout(doc, "Ponto de atenção", "Antes de produção, troque senhas iniciais, mantenha DATA_DIR fora do repositório, use HTTPS, revise CSP quando adicionar serviços externos e considere CSRF token para formulários administrativos.", "FFF7DB", "7A5A00")

    add_h1(doc, "7. Como Rodar e Publicar")
    add_h2(doc, "Local")
    add_numbers(doc, [
        "Instale Node.js 20 ou superior.",
        "Na pasta do projeto, execute npm install.",
        "Execute npm start.",
        "Acesse http://localhost:3000.",
        "Rode npm run check antes de publicar alterações."
    ])
    add_h2(doc, "Render")
    add_bullets(doc, [
        "Build Command: npm install.",
        "Start Command: npm start.",
        "Variável recomendada: NODE_ENV=production.",
        "Configure disco persistente para DATA_DIR, evitando perda dos JSON e uploads em redeploy.",
        "Conecte o serviço ao repositório GitHub e publique a partir do branch principal."
    ])

    add_h1(doc, "8. Manutenção e Evolução")
    add_bullets(doc, [
        "Executar backup geral antes de alterações grandes de dados.",
        "Não versionar data/, uploads ou arquivos com informações sensíveis.",
        "Testar login admin, login usuário comum, filtros, criação de lição, PDF/ebook e trilhas depois de mudanças em app.js, auth.js ou server.js.",
        "Evolução recomendada: migrar JSON para PostgreSQL, usar storage dedicado para uploads, integrar SMTP e WhatsApp Cloud API, adicionar testes automatizados e melhorar auditoria."
    ])

    doc.save(OUT_PROGRAMADOR)


def build_user_manual():
    doc = setup_doc("Manual do Usuário")
    add_footer(doc, "Manual do Usuário")
    add_cover(
        doc,
        "Manual do Usuário",
        "Guia prático de uso do sistema Raízes Kids",
        "Público: administradores, líderes, voluntários, pais e usuários cadastrados",
        GREEN,
    )

    add_h1(doc, "1. O que é o Raízes Kids")
    add_para(doc, "O Raízes Kids é um portal para organizar e acessar materiais de ensino bíblico infantil. Nele, o usuário encontra lições por idade, trilhas em vídeo, cultos em família, treinamentos, materiais de EBF e novidades do ministério.")
    add_callout(doc, "Ideia principal", "O sistema foi criado para reduzir o tempo de preparação e deixar o ensino mais organizado, bonito e acessível em celular, tablet, computador e Smart TV.", "F1FCEB", GREEN)

    add_h1(doc, "2. Acesso ao Sistema")
    add_h2(doc, "Entrar")
    add_numbers(doc, [
        "Clique em Entrar / Cadastrar.",
        "Informe usuário/CPF e senha.",
        "Clique em Entrar.",
        "Se o acesso estiver aprovado, o conteúdo será liberado conforme sua categoria."
    ])
    add_h2(doc, "Cadastrar novo usuário")
    add_numbers(doc, [
        "Abra a tela de cadastro.",
        "Preencha Nome Completo, CPF, Email, Endereço, Igreja e demais dados solicitados.",
        "Crie uma senha de 6 dígitos e confirme.",
        "Envie o cadastro e aguarde aprovação do administrador."
    ])
    add_h2(doc, "Redefinir senha")
    add_para(doc, "Na tela de login, use a opção de redefinição de senha. O administrador também pode criar uma nova senha para o usuário pela área de Gerenciamento.")

    add_h1(doc, "3. Tipos de Acesso")
    add_table(doc, ["Perfil/Categoria", "O que normalmente permite"], [
        ["Visitante", "Visualiza a vitrine e chamadas, mas encontra conteúdos bloqueados."],
        ["Usuário Simples", "Acesso básico a conteúdos liberados para acompanhamento familiar."],
        ["Teste", "Acesso temporário ou demonstrativo."],
        ["Líderes", "Acesso a lições e trilhas voltadas ao preparo de aulas."],
        ["Prime", "Acesso amplo, incluindo treinamentos e EBF quando liberados."],
        ["Administrador", "Gerencia conteúdos, usuários, novidades, contatos, backups e trilhas."],
        ["Administrador master", "Além das funções de admin, acessa Comunicação contínua."],
    ], [2600, 6760])

    add_h1(doc, "4. Tela Inicial")
    add_bullets(doc, [
        "O banner principal apresenta a proposta do ministério e botões para conhecer o catálogo ou escolher faixa etária.",
        "Os cards de faixa etária direcionam a navegação para o grupo escolhido.",
        "A Central de Novidades mostra conteúdos recém-publicados ou avisos importantes.",
        "O menu superior organiza conteúdos por Famílias, Líderes, Coordenadores e Administração."
    ])

    add_h1(doc, "5. Catálogo de Lições")
    add_h2(doc, "Como localizar uma lição")
    add_numbers(doc, [
        "Acesse Lições Bíblicas.",
        "Use Buscar para digitar tema, título ou versículo.",
        "Filtre por Categoria, Idade, Testamento, Data especial ou Mês de inclusão.",
        "Clique em uma lição para abrir a leitura completa."
    ])
    add_h2(doc, "Como estudar a lição")
    add_para(doc, "A página da lição apresenta o título, categoria, faixa etária, versículo e as partes da aula: objetivos, recepção, quebra-gelo, oração inicial, louvor/oferta, lição bíblica, aplicação prática, memorização, atividade, oração final e lanche quando cadastrados.")
    add_h2(doc, "Exportar PDF")
    add_para(doc, "Na visualização da lição, use o botão Exportar PDF. O navegador abrirá a impressão. Escolha Salvar como PDF. O conteúdo usa uma marca autoral e layout próprio de impressão.")

    add_h1(doc, "6. Trilhas com Vídeos")
    add_bullets(doc, [
        "Acesse Trilhas para ver vídeos organizados em formato de plataforma visual.",
        "Use os filtros superiores para ajustar por idade, categoria e busca.",
        "Clique em Assistir para abrir o player integrado na própria página.",
        "Quando necessário, use Abrir no YouTube para assistir diretamente no YouTube.",
        "As prateleiras agrupam vídeos por categoria e destaques como Em Alta."
    ])

    add_h1(doc, "7. Outros Conteúdos")
    add_table(doc, ["Área", "Como usar"], [
        ["Culto em Família", "Selecione um culto para visualizar princípio, texto bíblico, devocional, oração e atividade para casa."],
        ["Treinamentos", "Acesse materiais de capacitação, vídeos, orientações e anexos quando seu perfil permitir."],
        ["EBF Completa", "Consulte materiais de Escola Bíblica de Férias, programação, orientações, anexos e artes cadastradas."],
        ["Perfil", "Edite dados pessoais, telefone, endereço, igreja e cidade da igreja."],
        ["Contato", "Veja email, WhatsApp, Instagram e informações do ministério."],
    ], [2300, 7060])

    add_h1(doc, "8. Área de Gerenciamento")
    add_para(doc, "A área de Gerenciamento aparece apenas para administradores. Ela concentra cadastros, usuários, backups e configurações.")
    add_h2(doc, "Gerenciar lições")
    add_numbers(doc, [
        "Abra Gerenciamento e escolha a aba Lições.",
        "Clique em Nova lição para começar do zero ou selecione uma lição existente.",
        "Preencha título, categoria, idade, versículo, imagem do card, imagem da atividade de colorir e seções da aula.",
        "Clique em Salvar e aguarde a mensagem de confirmação.",
        "Use Limpar para esvaziar o formulário, Duplicar lição para reaproveitar uma base ou Excluir para remover."
    ])
    add_h2(doc, "Exportar ebook")
    add_para(doc, "No Gerenciamento, ajuste os filtros superiores e clique em Exportar Ebook. O ebook incluirá todas as lições ativas no filtro atual.")
    add_h2(doc, "Gerenciar trilhas")
    add_numbers(doc, [
        "Abra a aba Gerenciar trilhas.",
        "Informe título, link do YouTube, categoria, idade, lição relacionada, observação, playlist e temporada.",
        "Marque Destaque principal ou Em Alta quando quiser evidenciar o vídeo.",
        "Clique em Salvar vídeo e confira a confirmação."
    ])
    add_h2(doc, "Novidades")
    add_para(doc, "Na aba Novidades, cadastre chamadas com título, tipo, destino, botão, data de publicação, expiração, resumo, status ativa e destaque. Essas chamadas aparecem no sino e na área de novidades da home.")

    add_h1(doc, "9. Gestão de Usuários")
    add_bullets(doc, [
        "A aba Usuários lista cadastros aprovados, pendentes e inativos.",
        "O administrador pode aprovar, desativar, reativar, renovar licença, redefinir senha e alterar categoria de acesso.",
        "O botão Exportar Excel gera uma planilha CSV com os dados dos usuários.",
        "A aba Acesso mostra registros de uso, IP, dispositivo e página acessada; também permite exportar ou apagar logs."
    ])

    add_h1(doc, "10. Comunicação e Contato")
    add_h2(doc, "Comunicação contínua")
    add_para(doc, "Disponível apenas para o administrador master. Permite escolher canal, grupo de acesso, status dos usuários, assunto e mensagem. Ao criar campanha, o sistema gera links de email e WhatsApp para envio manual controlado.")
    add_h2(doc, "Sobre e contato")
    add_para(doc, "Na aba Contato do gerenciamento, o administrador edita o texto Sobre, email, WhatsApp, Instagram e site exibidos no rodapé.")

    add_h1(doc, "11. Backup")
    add_para(doc, "O botão Backup Geral baixa um arquivo compactado com os dados do sistema. Use antes de alterações grandes, importações ou manutenção em produção.")

    add_h1(doc, "12. Boas Práticas")
    add_bullets(doc, [
        "Ao cadastrar lições, mantenha títulos claros e categorias consistentes.",
        "Use imagens leves para melhorar desempenho em celulares.",
        "Antes de excluir, confirme se o conteúdo não será usado em aula ou ebook.",
        "Depois de salvar, sempre confira a mensagem de confirmação.",
        "Use filtros antes de exportar ebook para não imprimir conteúdos desnecessários.",
        "Mantenha dados de contato atualizados para facilitar suporte e comunicação."
    ])

    add_h1(doc, "13. Dúvidas Frequentes")
    add_table(doc, ["Situação", "O que fazer"], [
        ["Não consigo entrar.", "Verifique usuário e senha. Se necessário, solicite redefinição de senha."],
        ["Meu cadastro não libera conteúdo.", "Aguarde aprovação do administrador."],
        ["A lição não apareceu depois de salvar.", "Confira a mensagem de confirmação e atualize a página. Verifique também os filtros ativos."],
        ["PDF saiu estranho.", "Use Salvar como PDF no navegador e aguarde imagens carregarem antes de confirmar."],
        ["Vídeo não toca.", "Tente abrir no YouTube. Alguns vídeos bloqueiam incorporação pelo proprietário."],
        ["Não vejo Gerenciamento.", "A aba aparece somente para usuários administradores."],
    ], [2600, 6760])

    doc.save(OUT_USUARIO)


if __name__ == "__main__":
    build_programmer_manual()
    build_user_manual()
    print(OUT_PROGRAMADOR)
    print(OUT_USUARIO)
